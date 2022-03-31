import subprocess

import docx.table
import pandas as pd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml import ns
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.table import _Cell
from docx2pdf import convert


class BasicReport:
    def __init__(self):
        pass

    @staticmethod
    def column_text_change(data: pd.DataFrame,
                           table: docx.table.Table,
                           size: int,
                           bold: bool = False):
        for i in range(len(data.columns)):
            run = table.cell(0, i).paragraphs[0].runs[0]
            run.font.size = Pt(size)
            run.font.bold = bold

    @staticmethod
    def cell_text_change(data: pd.DataFrame,
                         table: docx.table.Table,
                         size: int,
                         bold: bool = False):
        for i in range(data.shape[0]):
            for j in range(data.shape[1]):
                run = table.cell(i+1, j).paragraphs[0].runs[0]
                run.font.size = Pt(size)
                run.font.bold = bold

    @staticmethod
    def delete_columns(table, columns):
        # sort columns descending
        columns.sort(reverse=True)

        grid = table._tbl.find("w:tblGrid", table._tbl.nsmap)
        for ci in columns:
            for cell in table.column_cells(ci):
                cell._tc.getparent().remove(cell._tc)

            # Delete column reference.
            col_elem = grid[ci]
            grid.remove(col_elem)

    @staticmethod
    def generate_pdf_windows(doc_path: str, out_path: str):
        """Generate pdf file for windows system"""
        convert(doc_path, out_path)

    @staticmethod
    def generate_pdf_Linux(doc_path, out_path):
        """Generate pdf file for windows system"""
        subprocess.call(['soffice',
                         # '--headless',
                         '--convert-to',
                         'pdf',
                         '--outdir',
                         out_path,
                         doc_path])
        return doc_path

    @staticmethod
    def set_row_height(row, height):
        trPr = row.tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), str(height))
        trPr.append(trHeight)

    @staticmethod
    def set_vertical_cell_direction(cell: _Cell, direction: str):
        # direction: tbRl -- top to bottom, btLr -- bottom to top
        assert direction in ("tbRl", "btLr")
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        textDirection = OxmlElement('w:textDirection')
        textDirection.set(qn('w:val'), direction)  # btLr tbRl
        tcPr.append(textDirection)

    @staticmethod
    def create_element(name):
        return OxmlElement(name)

    @staticmethod
    def create_attribute(element, name, value):
        element.set(ns.qn(name), value)

    def create_document(self, header):
        document = Document()
        section = document.sections[-1]
        section.left_martin = Inches(0.1)
        paragraph_format = document.styles['Normal'].paragraph_format
        paragraph_format.space_before = 0
        paragraph_format.space_after = 0

        document.add_paragraph().add_run(header).bold = True
        document.add_paragraph(" ")

        section.footer.paragraphs[0].text = header
        section.footer.add_paragraph()
        self.add_page_number(section.footer.paragraphs[1].add_run())
        section.footer.paragraphs[1].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        return document

    def add_page_number(self, run):
        fldChar1 = self.create_element('w:fldChar')
        self.create_attribute(fldChar1, 'w:fldCharType', 'begin')

        instrText = self.create_element('w:instrText')
        self.create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = "PAGE"

        fldChar2 = self.create_element('w:fldChar')
        self.create_attribute(fldChar2, 'w:fldCharType', 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
