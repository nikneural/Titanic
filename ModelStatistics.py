from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches

from BasicReport import BasicReport


class ModelStatistics(BasicReport):
    def __init__(self, report_number: int, data: str):
        super().__init__()
        self.report_number = report_number
        self.data = data

    def creating_word_file(self):
        document = Document()
        section = document.sections[-1]
        section.left_martin = Inches(0.1)
        paragraph_format = document.styles['Normal'].paragraph_format
        paragraph_format.space_before = 0
        paragraph_format.space_after = 0

        header = f"Отчет №{self.report_number} от {self.data}. Анализ работы модели BTC"
        section.footer.paragraphs[0].text = header
        section.footer.add_paragraph()
        self.add_page_number(section.footer.paragraphs[1].add_run())
        section.footer.paragraphs[1].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        document.add_paragraph().add_run(header).bold = True
        document.add_paragraph("В отчете проведен анализ работы модели с отключением торговли по состоянию депозита "
                               "отдельно для следующих периодов:")
        document.add_paragraph('\tTrain. Обучающий период с 2017.09 по 2021.01')
        document.add_paragraph('\tTest. Тестовый период с 2021.01 по 2021.12')
        document.add_paragraph('\tTest - 12.2021')
        document.add_paragraph('\tTest - 01.2022')
        document.add_paragraph('\tTest - 02.2022')

        document.add_paragraph("Анализ проводится отдельно для сделок, серий сделок и помесячных периодов. Если "
                               "параметры торговли за последний период будут выходить за границы распределений "
                               "параметров торговли за предыдущие периоды, то на них нельзя будет опираться при "
                               "оценке статистики работы модели.")
        document.add_paragraph("\n")
        document.add_paragraph().add_run("Анализ для серий сделок").bold = True
        document.add_paragraph("На графике представлены относительные распределения по следующим параметрам:")
        document.add_paragraph('Series_sum - сумма серии сделок')
        document.add_paragraph('D_prc - изменение курса за период серии в процентах')
        document.add_paragraph('Nstd - нормированное стандартное отклонение курса за период серии')
        document.add_paragraph('Max_prc - максимальное изменение цены вверх в процентах от начальной цены за период '
                               'серии')
        document.add_paragraph('Min_prc - просадка цены в процентах от начальной цены за период серии')
        document.add_paragraph().add_run('BTC').bold = True
        # картинка
        document.add_page_break()
        document.add_paragraph().add_run('Анализ для сделок').bold = True
        document.add_paragraph('На графике представлены относительные распределения по следующим параметрам:')
        document.add_paragraph('D_prc - прибыль сделки')
        document.add_paragraph('Nstd - нормированное стандартное отклонение курса за период сделки')
        document.add_paragraph('Max_prc - максимальное изменение цены вверх в процентах от начальной цены за период '
                               'сделки')
        document.add_paragraph('Min_prc - просадка цены в процентах от начальной цены за период сделки')
        document.add_paragraph().add_run('BTC').bold = True
        # картинка
        document.add_page_break()
        document.add_paragraph().add_run('Анализ для помесячных периодов').bold = True
        document.add_paragraph("На графике представлены относительные распределения по следующим параметрам:")
        document.add_paragraph('btc_move - изменение курса btc')
        document.add_paragraph('btc_min_prc - просадка цены в процентах от начальной цены за период')
        document.add_paragraph(
            'btc_max_prc - максимальное изменение цены вверх в процентах от начальной цены за период')
        document.add_paragraph('btc_nstd - нормированное стандартное отклонение курса за период')
        document.add_paragraph('sum - прибыль за период')
        document.add_paragraph('markup% - процент прибыли от авторазметки')
        document.add_paragraph().add_run('BTC').bold = True
        # картинка

        document.save('ModelStatistics.docx')
